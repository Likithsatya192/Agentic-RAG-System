import os
import logging
from typing import List, Dict, Any
from pathlib import Path

import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentLoader:
    """Handles loading and processing of various document types"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
        )
        
    def load_pdf(self, file_path: str) -> List[str]:
        """Load and extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return [text]
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            return []
    
    def load_docx(self, file_path: str) -> List[str]:
        """Load and extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return [text]
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            return []
    
    def load_txt(self, file_path: str) -> List[str]:
        """Load and extract text from TXT files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return [file.read()]
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            return []
    
    def load_documents_from_directory(self, directory: str) -> List[LangChainDocument]:
        """Load all supported documents from a directory"""
        documents = []
        directory_path = Path(directory)
        
        if not directory_path.exists():
            logger.error(f"Directory {directory} does not exist")
            return documents
        
        for file_path in directory_path.rglob("*"):
            if file_path.is_file():
                file_extension = file_path.suffix.lower()
                texts = []
                
                if file_extension == '.pdf':
                    texts = self.load_pdf(str(file_path))
                elif file_extension == '.docx':
                    texts = self.load_docx(str(file_path))
                elif file_extension == '.txt':
                    texts = self.load_txt(str(file_path))
                
                for text in texts:
                    if text.strip():  # Only add non-empty texts
                        documents.append(
                            LangChainDocument(
                                page_content=text,
                                metadata={"source": str(file_path)}
                            )
                        )
        
        logger.info(f"Loaded {len(documents)} documents from {directory}")
        return documents
    
    def create_vector_store(self, documents: List[LangChainDocument]) -> FAISS:
        """Create a vector store from documents"""
        if not documents:
            logger.warning("No documents provided for vector store creation")
            return None
        
        # Split documents into chunks
        split_docs = self.text_splitter.split_documents(documents)
        
        # Create embeddings
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        # Create vector store
        vector_store = FAISS.from_documents(split_docs, embeddings)
        logger.info(f"Created vector store with {len(split_docs)} chunks")
        
        return vector_store